import os
from docx import Document
from deep_translator import GoogleTranslator

# Folders
NEPALI_DOCX_FOLDER = os.path.join(os.path.dirname(__file__), "nepali_transcript")
ENGLISH_DOCX_FOLDER = os.path.join(os.path.dirname(__file__), "nepali_translate")

# Ensure output folder exists
os.makedirs(ENGLISH_DOCX_FOLDER, exist_ok=True)

def translate_paragraph(text):
    try:
        return GoogleTranslator(source='auto', target='en').translate(text)
    except Exception as e:
        print(f"⚠️ Translation failed for: {text[:40]}... Error: {e}")
        return text  # fallback to original

def translate_docx_file(input_path, output_path):
    doc = Document(input_path)
    translated_doc = Document()

    for para in doc.paragraphs:
        text = para.text.strip()
        translated_text = translate_paragraph(text) if text else ""
        translated_doc.add_paragraph(translated_text)

    translated_doc.save(output_path)

def main():
    files = [f for f in os.listdir(NEPALI_DOCX_FOLDER) if f.endswith('.docx')]
    print(f"Found {len(files)} Nepali Word files to translate.")

    for file in files:
        input_path = os.path.join(NEPALI_DOCX_FOLDER, file)
        output_path = os.path.join(ENGLISH_DOCX_FOLDER, file.replace(".docx", "_en.docx"))
        try:
            translate_docx_file(input_path, output_path)
            print(f"✅ Translated and saved: {output_path}")
        except Exception as e:
            print(f"❌ Failed to translate {file}: {e}")

if __name__ == "__main__":
    main()
